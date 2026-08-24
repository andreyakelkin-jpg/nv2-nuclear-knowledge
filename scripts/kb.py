#!/usr/bin/env python3
"""Простой и проверяемый конвейер нормативной базы знаний.

Скрипт не делает юридических выводов и не подменяет Архивария эвристиками.
Он сохраняет оригинал, применяет структурированное решение Архивария,
перестраивает реестры и откатывает всю операцию при ошибке целостности.
"""
from __future__ import annotations

import argparse
import hashlib
import re
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import yaml

from kb_root import CONFIG_PATH, SUPPORTED_SCHEMA_VERSION, resolve_kb_root
from model_router import (
    apply_quality_gate,
    check_route,
    claim_side_effects,
    escalate_route,
    routing_status,
    set_routing_enabled,
    start_route,
)
from reference_resolver import (
    canonical_identifier,
    identifier_from_id,
    queue_key,
    queue_score,
    reference_role,
    synchronize_references,
)
from reference_store import (
    REFERENCE_SCHEMA_VERSION,
    attach_reference_graph,
    externalize_all,
    load_references,
    reference_path,
    write_references,
)
from retrieval import archive_context, build_retrieval_indexes, fetch_document, print_json, search_documents


ROOT = resolve_kb_root()
META = ROOT / "meta"
VALID_REFERENCE_STATUSES = {
    "в_базе", "отсутствует", "устарел_есть_замена",
    "устарел_нет_замены", "устарел_специфика",
}
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
SAFE_ID = re.compile(r"^[a-z0-9][a-z0-9-]*$")
MD_LINK = re.compile(r"\]\(([^)]+\.md(?:#[^)]+)?)\)")


def now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def read_yaml(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    loaded = yaml.safe_load(path.read_text(encoding="utf-8"))
    return loaded or {}


def write_yaml(path: Path, data: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(yaml.safe_dump(data, allow_unicode=True, sort_keys=False, width=100), encoding="utf-8")


def relative(path: Path) -> str:
    return path.resolve().relative_to(ROOT.resolve()).as_posix()


def markdown_with_front_matter(metadata: dict[str, Any], body: str) -> str:
    return "---\n" + yaml.safe_dump(metadata, allow_unicode=True, sort_keys=False, width=100) + "---\n" + body


def front_matter(path: Path) -> tuple[dict[str, Any], str]:
    text = path.read_text(encoding="utf-8")
    if not text.startswith("---\n"):
        raise ValueError("нет YAML-фронтматтера")
    end = text.find("\n---\n", 4)
    if end < 0:
        raise ValueError("не закрыт YAML-фронтматтер")
    metadata = yaml.safe_load(text[4:end]) or {}
    if not isinstance(metadata, dict):
        raise ValueError("YAML-фронтматтер должен быть объектом")
    return metadata, text[end + 5:]


def document_files() -> list[Path]:
    docs = ROOT / "docs"
    return sorted(path for path in docs.rglob("*.md") if "_templates" not in path.parts)


def reference_files() -> list[Path]:
    return sorted((ROOT / "relations" / "references").glob("*.yaml"))


def expected_reference_files() -> list[Path]:
    paths: list[Path] = []
    for path in document_files():
        metadata, _ = front_matter(path)
        doc_id = str(metadata.get("id", ""))
        if SAFE_ID.fullmatch(doc_id):
            paths.append(reference_path(ROOT, doc_id))
    return paths


def known_categories() -> set[str]:
    return {item["id"] for item in read_yaml(META / "categories.yaml").get("categories", [])}


def type_directory(document_type: str) -> str:
    aliases = {
        "НП": "np", "ФНП": "fnp", "ГОСТ": "gost", "ГОСТ Р": "gost-r",
        "ТУ": "tu", "ОТТ": "ott", "СТО": "sto", "РД": "rd",
    }
    normalized = re.sub(r"\s+", " ", document_type.strip().upper())
    if normalized in aliases:
        return aliases[normalized]
    slug = re.sub(r"[^a-z0-9-]+", "-", document_type.lower()).strip("-")
    return slug or "other"


def extract_text(source: Path) -> tuple[str, str]:
    suffix = source.suffix.lower()
    if suffix in {".txt", ".md"}:
        return source.read_text(encoding="utf-8", errors="replace"), "plain-text"
    if suffix == ".docx":
        from docx import Document

        document = Document(source)
        parts = [item.text for item in document.paragraphs if item.text.strip()]
        for number, table in enumerate(document.tables, 1):
            parts.append(f"\n[ТАБЛИЦА {number}]")
            parts.extend(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)
        return "\n".join(parts), "python-docx"
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
            pages = PdfReader(str(source)).pages
            return "".join(
                f"\n\n===== СТРАНИЦА {number} =====\n{page.extract_text() or ''}"
                for number, page in enumerate(pages, 1)
            ), "pypdf"
        except Exception:
            import pdfplumber
            with pdfplumber.open(source) as pdf:
                return "".join(
                    f"\n\n===== СТРАНИЦА {number} =====\n{page.extract_text() or ''}"
                    for number, page in enumerate(pdf.pages, 1)
                ), "pdfplumber"
    raise ValueError("Допустимы только PDF, DOCX, TXT и MD")


def existing_hashes() -> set[str]:
    hashes: set[str] = set()
    for path in document_files():
        try:
            metadata, _ = front_matter(path)
            digest = metadata.get("source", {}).get("sha256")
            if digest:
                hashes.add(str(digest))
        except ValueError:
            pass
    return hashes


def staged_hashes() -> set[str]:
    hashes: set[str] = set()
    for manifest_path in (ROOT / "staging").glob("*/manifest.yaml"):
        digest = read_yaml(manifest_path).get("source_sha256")
        if digest:
            hashes.add(str(digest))
    return hashes


def stage_document(source: Path) -> str:
    source = source.resolve()
    if not source.is_file():
        raise FileNotFoundError(f"Файл не найден: {source}")
    if source.suffix.lower() not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Неподдерживаемый формат {source.suffix}")
    digest = hashlib.sha256(source.read_bytes()).hexdigest()
    if digest in existing_hashes():
        return "ALREADY_ARCHIVED"
    if digest in staged_hashes():
        return "ALREADY_STAGED"
    stage_id = f"{datetime.now():%Y%m%d-%H%M%S}-{digest[:10]}"
    stage_dir = ROOT / "staging" / stage_id
    stage_dir.mkdir(parents=True)
    staged_source = stage_dir / f"source{source.suffix.lower()}"
    shutil.copy2(source, staged_source)
    text, method = extract_text(staged_source)
    (stage_dir / "extracted.txt").write_text(text, encoding="utf-8")
    write_yaml(stage_dir / "manifest.yaml", {
        "stage_id": stage_id,
        "state": "waiting_for_ai_analysis",
        "received_at": now(),
        "source_name": source.name,
        "source_sha256": digest,
        "source_extension": source.suffix.lower(),
        "staged_source": relative(staged_source),
        "extracted_text": relative(stage_dir / "extracted.txt"),
        "extraction_method": method,
        "characters_extracted": len(text),
        "requires_visual_review": len(text.strip()) < 500,
        "warning": "Мало извлечённого текста: требуется OCR и визуальная проверка." if len(text.strip()) < 500 else None,
    })
    return stage_id


def command_stage(args: argparse.Namespace) -> None:
    result = stage_document(Path(args.source))
    if result.startswith("ALREADY"):
        print("Документ уже есть в базе." if result == "ALREADY_ARCHIVED" else "Документ уже ожидает анализа.")
        return
    print(f"Готово к анализу: staging/{result}/manifest.yaml")
    print("Следующий шаг: прикрепите исходник в чат или передайте Архиварию manifest.yaml и extracted.txt.")


def command_intake(_: argparse.Namespace) -> None:
    files = [path for path in sorted((ROOT / "inbox").iterdir()) if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS]
    if not files:
        print("Папка inbox пуста.")
        return
    for source in files:
        try:
            result = stage_document(source)
            print(f"{source.name}: {result}")
        except Exception as error:
            print(f"{source.name}: ОШИБКА — {error}")


def snapshot(paths: list[Path]) -> dict[Path, bytes | None]:
    return {path: path.read_bytes() if path.exists() else None for path in paths}


def restore(data: dict[Path, bytes | None]) -> None:
    for path, content in data.items():
        if content is None:
            if path.exists():
                path.unlink()
        else:
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_bytes(content)


def build_indexes() -> dict[str, int]:
    documents: list[dict[str, Any]] = []
    references: list[dict[str, Any]] = []
    materials: dict[str, dict[str, Any]] = {}
    impacts: dict[str, dict[str, Any]] = {}
    previous_queue = {
        queue_key({"cited_as": str(item.get("document_label", ""))}): item
        for item in read_yaml(META / "addition-queue.yaml").get("queue", [])
    }
    generated_queue: dict[str, dict[str, Any]] = {}
    generated_replacements: dict[str, dict[str, Any]] = {}
    indexed_at = now()
    for path in document_files():
        metadata, _ = front_matter(path)
        doc_id = str(metadata["id"])
        document_references = load_references(ROOT, metadata, allow_legacy=True)
        source = metadata.get("source", {})
        exact, family = canonical_identifier(str(metadata.get("short_title") or metadata.get("title") or doc_id))
        fallback_exact, fallback_family = identifier_from_id(doc_id)
        relations = metadata.get("relations", {})
        verification = metadata.get("verification", {})
        documents.append({
            "id": doc_id, "path": relative(path), "status": metadata.get("status", "требует_проверки"),
            "lifecycle_stage": metadata.get("lifecycle", {}).get("stage", "requires_expert_review"),
            "type": metadata.get("type"), "title": metadata.get("title"), "short_title": metadata.get("short_title"),
            "issuer": metadata.get("issuer"), "effective_from": metadata.get("effective_from"),
            "revision_date": metadata.get("revision_date"),
            "categories": metadata.get("category", []), "source_available": bool(source.get("original_file")),
            "applies_to": metadata.get("applies_to", []),
            "normalized_path": source.get("normalized_file"),
            "source_sha256": source.get("sha256"),
            "verification": {
                "legal_status": verification.get("legal_status"),
                "legal_status_checked_on": verification.get("legal_status_checked_on"),
                "legal_status_source": verification.get("legal_status_source"),
                "reviewed_by": verification.get("reviewed_by"),
                "reviewed_on": verification.get("reviewed_on"),
            },
            "replaces": relations.get("replaces", []), "replaced_by": relations.get("replaced_by", []),
            "related_documents": relations.get("related_documents", []),
            "reference_graph": metadata.get("reference_graph", {"count": len(document_references)}),
            "canonical_exact": exact or fallback_exact, "canonical_family": family or fallback_family,
            "indexed_at": indexed_at,
        })
        for number, reference in enumerate(document_references, 1):
            item = dict(reference)
            item.update({"id": f"ref-{doc_id}-{number:03d}", "source_document": doc_id})
            references.append(item)
            for linked_id, relation in ((item.get("target_document"), "прямая_ссылка"), (item.get("replacement_document"), "замена")):
                if linked_id:
                    entry = impacts.setdefault(str(linked_id), {"document_id": str(linked_id), "affects": []})
                    entry["affects"].append({"document_id": doc_id, "reference_id": item["id"], "relation": relation})
            if item.get("status") in {"отсутствует", "устарел_нет_замены"}:
                label = str(item.get("cited_as", "неидентифицированный документ"))
                key = queue_key(item)
                prior = previous_queue.get(key, {})
                entry = generated_queue.setdefault(key, {
                    "id": prior.get("id", f"add-{hashlib.sha1(key.encode()).hexdigest()[:12]}"),
                    "canonical_key": key,
                    "variants": set(),
                    "cited_by": set(),
                    "reference_roles": set(),
                    "reference_statuses": set(),
                    "reasons": set(),
                    "status": prior.get("status", "requested"),
                    "user_note": prior.get("user_note"),
                    "priority_locked": bool(prior.get("priority_locked", False)),
                    "locked_priority": prior.get("priority"),
                })
                entry["variants"].add(label)
                entry["cited_by"].add(doc_id)
                entry["reference_roles"].add(item.get("reference_role") or reference_role(item))
                entry["reference_statuses"].add(str(item.get("status")))
                entry["reasons"].add(item.get("action") or item.get("basis") or "Указан в нормативном документе")
            if item.get("status") == "устарел_есть_замена" and item.get("replacement_document"):
                replacement_key = queue_key(item)
                replacement = generated_replacements.setdefault(replacement_key, {
                    "old_document": item.get("cited_as"),
                    "replacement_document": item.get("replacement_document"),
                    "status": "requires_expert_review",
                    "basis": set(),
                    "cited_by": set(),
                })
                if item.get("basis"):
                    replacement["basis"].add(str(item["basis"]))
                replacement["cited_by"].add(doc_id)
        for material in metadata.get("materials", []):
            name = material.get("name") if isinstance(material, dict) else str(material)
            if name:
                materials.setdefault(name.lower(), {"name": name, "mentioned_in": []})["mentioned_in"].append(doc_id)
    documents.sort(key=lambda value: value["id"])
    references.sort(key=lambda value: value["id"])
    queue_output: list[dict[str, Any]] = []
    for entry in generated_queue.values():
        variants = sorted(entry.pop("variants"), key=lambda value: (len(value), value))
        cited_by = sorted(entry.pop("cited_by"))
        roles = sorted(entry.pop("reference_roles"))
        statuses = set(entry.pop("reference_statuses"))
        reasons = sorted(entry.pop("reasons"))
        score, calculated_priority = queue_score(set(roles), len(cited_by), statuses)
        locked_priority = entry.pop("locked_priority")
        priority = locked_priority if entry.get("priority_locked") and locked_priority else calculated_priority
        entry.update({
            "document_label": variants[0],
            "variants": variants,
            "cited_by": cited_by,
            "reference_roles": roles,
            "reason": reasons[0] if reasons else "Указан в нормативном документе",
            "score": score,
            "priority": priority,
        })
        queue_output.append(entry)
    priority_order = {"high": 0, "medium": 1, "low": 2}
    queue_output.sort(key=lambda value: (priority_order.get(value["priority"], 9), -value["score"], value["document_label"]))
    replacements_output = []
    for key, entry in sorted(generated_replacements.items()):
        entry["canonical_key"] = key
        entry["basis"] = sorted(entry["basis"])
        entry["cited_by"] = sorted(entry["cited_by"])
        replacements_output.append(entry)
    write_yaml(META / "documents.yaml", {"documents": documents})
    write_yaml(META / "cross-references.yaml", {"references": references})
    write_yaml(META / "materials.yaml", {"materials": sorted(materials.values(), key=lambda value: value["name"])})
    write_yaml(META / "addition-queue.yaml", {"queue": queue_output})
    write_yaml(META / "replacements.yaml", {"replacements": replacements_output})
    write_yaml(META / "impact-index.yaml", {"impacts": sorted(impacts.values(), key=lambda value: value["document_id"])})
    retrieval_summary = build_retrieval_indexes(ROOT, documents, indexed_at)
    write_yaml(META / "corpus-manifest.yaml", {
        "schema_version": 2, "last_indexed_at": indexed_at, "document_count": len(documents),
        "reference_count": len(references),
        "reference_storage_schema": 1,
        "retrieval_index": retrieval_summary,
        "context_policy": "В LLM передаются карточки, реестры и релевантные фрагменты, а не весь корпус.",
    })
    return {
        "documents": len(documents), "references": len(references), "queue": len(queue_output),
        "clauses": retrieval_summary["clauses"], "pages": retrieval_summary["pages"],
    }


def validate_base() -> tuple[list[str], list[str]]:
    errors: list[str] = []
    warnings: list[str] = []
    ids: set[str] = set()
    categories = known_categories()
    confidence_missing = 0
    for path in document_files():
        try:
            metadata, body = front_matter(path)
        except Exception as error:
            errors.append(f"{relative(path)}: {error}")
            continue
        doc_id = str(metadata.get("id", ""))
        if not SAFE_ID.fullmatch(doc_id): errors.append(f"{relative(path)}: некорректный id {doc_id!r}")
        if doc_id in ids: errors.append(f"дублирующийся id: {doc_id}")
        ids.add(doc_id)
        unknown = set(metadata.get("category", [])) - categories
        if unknown: errors.append(f"{relative(path)}: неизвестные категории {sorted(unknown)}")
        source = metadata.get("source", {})
        if not source.get("sha256"): warnings.append(f"{doc_id}: нет SHA-256 оригинала")
        lifecycle = metadata.get("lifecycle", {}).get("stage")
        if lifecycle == "approved_for_operational_use":
            verification = metadata.get("verification", {})
            if not verification.get("legal_status_source"):
                errors.append(f"{doc_id}: рабочее применение без источника проверки статуса")
            if not verification.get("reviewed_by"):
                errors.append(f"{doc_id}: рабочее применение без проверяющего эксперта")
        try:
            document_references = load_references(ROOT, metadata, allow_legacy=False)
        except Exception as error:
            errors.append(f"{doc_id}: граф ссылок — {error}")
            document_references = []
        for number, reference in enumerate(document_references, 1):
            status = reference.get("status")
            if status not in VALID_REFERENCE_STATUSES:
                errors.append(f"{doc_id}: ссылка {number}, неверный статус {status!r}")
            if status == "в_базе" and not reference.get("target_document"):
                errors.append(f"{doc_id}: ссылка {number} в_базе без target_document")
            if status == "устарел_есть_замена" and not reference.get("replacement_document"):
                errors.append(f"{doc_id}: ссылка {number} без replacement_document")
            if status == "устарел_специфика" and not reference.get("basis"):
                errors.append(f"{doc_id}: ссылка {number} без документированного основания")
            if reference.get("confidence") in {None, "not_assessed"}:
                confidence_missing += 1
        for link in MD_LINK.findall(body):
            candidate = (path.parent / link.split("#", 1)[0]).resolve()
            if ROOT.resolve() not in candidate.parents or not candidate.exists():
                errors.append(f"{doc_id}: битая Markdown-ссылка {link}")
    for path in document_files():
        metadata, _ = front_matter(path)
        try:
            document_references = load_references(ROOT, metadata, allow_legacy=False)
        except Exception:
            continue
        for reference in document_references:
            if reference.get("status") == "в_базе" and reference.get("target_document") not in ids:
                errors.append(f"{metadata.get('id')}: цель ссылки отсутствует: {reference.get('target_document')}")
    expected_graphs = {reference_path(ROOT, doc_id).resolve() for doc_id in ids if SAFE_ID.fullmatch(doc_id)}
    for orphan in sorted({path.resolve() for path in reference_files()} - expected_graphs):
        errors.append(f"Граф ссылок без карточки: {relative(orphan)}")
    if confidence_missing:
        warnings.append(f"Для {confidence_missing} ссылок уверенность ещё не оценена")
    return errors, warnings


def write_validation_report(errors: list[str], warnings: list[str]) -> None:
    write_yaml(ROOT / "reports" / "integrity-latest.yaml", {
        "checked_at": now(), "valid": not errors, "errors": errors, "warnings": warnings,
    })


def command_rebuild_index(_: argparse.Namespace) -> None:
    sync = synchronize_references(ROOT)
    summary = build_indexes()
    print(f"Синхронизировано ссылок: {sync['resolved_references']}; изменено карточек: {sync['changed_documents']}")
    print(
        f"Индексы обновлены: документов {summary['documents']}, ссылок {summary['references']}, "
        f"пунктов {summary['clauses']}, страниц {summary['pages']}, очередь {summary['queue']}"
    )


def command_sync(_: argparse.Namespace) -> None:
    tracked = [
        *document_files(), *reference_files(), *expected_reference_files(),
        *META.glob("*.yaml"), *META.glob("*.json"),
    ]
    before = snapshot(tracked)
    try:
        sync = synchronize_references(ROOT)
        summary = build_indexes()
        errors, warnings = validate_base()
        if errors:
            raise ValueError("; ".join(errors))
        write_validation_report(errors, warnings)
    except Exception:
        restore(before)
        raise
    print(f"Ссылки синхронизированы: разрешено {sync['resolved_references']}, исправлено {sync['repaired_references']}, карточек изменено {sync['changed_documents']}")
    print(f"Очередь пересобрана: {summary['queue']} позиций")
    if sync["ambiguous"]:
        print(f"Неоднозначных обозначений: {len(sync['ambiguous'])}; требуется экспертный анализ")


def command_validate(_: argparse.Namespace) -> None:
    errors, warnings = validate_base()
    write_validation_report(errors, warnings)
    if warnings: print("ПРЕДУПРЕЖДЕНИЯ:\n" + "\n".join(f"- {item}" for item in warnings))
    if errors:
        print("ОШИБКИ ЦЕЛОСТНОСТИ:\n" + "\n".join(f"- {item}" for item in errors))
        raise SystemExit(1)
    print("Проверка пройдена.")


def command_migrate_references(args: argparse.Namespace) -> None:
    tracked = [
        *document_files(), *reference_files(), *expected_reference_files(),
        *META.glob("*.yaml"), *META.glob("*.json"),
    ]
    if args.dry_run:
        summary = externalize_all(ROOT, write=False)
        print_json({key: value for key, value in summary.items() if key != "results"})
        return
    before = snapshot(tracked)
    try:
        migration = externalize_all(ROOT, write=True)
        sync = synchronize_references(ROOT)
        summary = build_indexes()
        errors, warnings = validate_base()
        if errors:
            raise ValueError("; ".join(errors))
        write_validation_report(errors, warnings)
    except Exception:
        restore(before)
        raise
    print_json({
        "migration": {key: value for key, value in migration.items() if key != "results"},
        "reference_sync": sync,
        "indexes": summary,
        "warnings": warnings,
    })


def command_apply(args: argparse.Namespace) -> None:
    decision_path = Path(args.decision).resolve()
    decision = read_yaml(decision_path)
    required = {"stage_id", "document_id", "document_type", "category", "markdown_file"}
    if missing := required - set(decision):
        raise ValueError(f"В decision.yaml отсутствуют поля: {', '.join(sorted(missing))}")
    doc_id, category = str(decision["document_id"]), str(decision["category"])
    if not SAFE_ID.fullmatch(doc_id): raise ValueError("document_id: только строчные латинские буквы, цифры и дефисы")
    stage_dir = ROOT / "staging" / str(decision["stage_id"])
    manifest = read_yaml(stage_dir / "manifest.yaml")
    if not manifest: raise ValueError("Не найдена стадия загрузки")
    markdown_path = (ROOT / str(decision["markdown_file"])).resolve()
    if not markdown_path.is_file() or ROOT.resolve() not in markdown_path.parents:
        raise ValueError("markdown_file должен находиться в репозитории")
    metadata, body = front_matter(markdown_path)
    if metadata.get("id") != doc_id: raise ValueError("id карточки не совпадает с document_id")
    if category not in metadata.get("category", []): raise ValueError("категория отсутствует во фронтматтере")
    legacy_references = metadata.pop("references", None)
    references_file = decision.get("references_file")
    if references_file:
        references_path = (ROOT / str(references_file)).resolve()
        if not references_path.is_file() or ROOT.resolve() not in references_path.parents:
            raise ValueError("references_file должен находиться в репозитории")
        references_payload = read_yaml(references_path)
        if references_payload.get("schema_version") != REFERENCE_SCHEMA_VERSION:
            raise ValueError("references_file содержит неподдерживаемую версию схемы")
        if references_payload.get("document_id") not in {None, doc_id}:
            raise ValueError("document_id в references_file не совпадает с decision.yaml")
        document_references = references_payload.get("references", [])
        if legacy_references is not None and legacy_references != document_references:
            raise ValueError("references в карточке расходятся с references_file")
    else:
        document_references = legacy_references or []
    if not isinstance(document_references, list) or any(not isinstance(item, dict) for item in document_references):
        raise ValueError("references_file должен содержать массив объектов references")
    for reference in document_references:
        if reference.get("status") not in VALID_REFERENCE_STATUSES:
            raise ValueError(f"Некорректный статус ссылки: {reference.get('status')}")
    new_category = decision.get("new_category")
    if category not in known_categories() and (not isinstance(new_category, dict) or new_category.get("id") != category):
        raise ValueError("Неизвестная категория без блока new_category")
    if new_category:
        for field in ("id", "title", "description"):
            if not new_category.get(field): raise ValueError(f"В new_category отсутствует {field}")
    source_files = list(stage_dir.glob("source.*"))
    if len(source_files) != 1: raise ValueError("В стадии должен быть один исходный файл")
    type_slug = type_directory(str(decision["document_type"]))
    destination_md = ROOT / "docs" / category / type_slug / f"{doc_id}.md"
    raw_path = ROOT / "raw" / category / type_slug / f"{doc_id}{source_files[0].suffix.lower()}"
    normalized_path = ROOT / "normalized" / f"{doc_id}.txt"
    destination_references = reference_path(ROOT, doc_id)
    if destination_md.exists() and not args.replace:
        raise FileExistsError(f"Карточка уже существует: {relative(destination_md)}. Для новой редакции используйте новый id.")
    tracked = [
        destination_md, raw_path, normalized_path, destination_references,
        *document_files(), *reference_files(), *expected_reference_files(), META / "impact-index.yaml",
        *META.glob("*.yaml"), *META.glob("*.json"), META / "clause-index.json", META / "search-index.json",
    ]
    before = snapshot(tracked)
    try:
        if new_category and category not in known_categories():
            categories = read_yaml(META / "categories.yaml")
            categories.setdefault("categories", []).append(new_category)
            write_yaml(META / "categories.yaml", categories)
        source = metadata.setdefault("source", {})
        source.update({"original_file": relative(raw_path), "normalized_file": relative(normalized_path),
                       "sha256": manifest["source_sha256"], "extraction_method": manifest["extraction_method"]})
        metadata.setdefault("lifecycle", {}).setdefault("stage", "requires_expert_review")
        metadata.setdefault("provenance", {}).update({"stage_id": manifest["stage_id"], "archived_at": now()})
        destination_md.parent.mkdir(parents=True, exist_ok=True)
        raw_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source_files[0], raw_path)
        shutil.copy2(stage_dir / "extracted.txt", normalized_path)
        metadata.pop("reference_graph", None)
        write_references(ROOT, doc_id, document_references)
        attach_reference_graph(ROOT, metadata, document_references)
        destination_md.write_text(markdown_with_front_matter(metadata, body), encoding="utf-8")
        sync = synchronize_references(ROOT)
        summary = build_indexes()
        errors, warnings = validate_base()
        if errors: raise ValueError("; ".join(errors))
        write_validation_report(errors, warnings)
    except Exception:
        restore(before)
        raise
    report = {"archived_at": now(), "document_id": doc_id, "stage_id": manifest["stage_id"],
              "summary": summary, "reference_sync": sync, "warnings": warnings,
              "lifecycle_stage": metadata["lifecycle"]["stage"]}
    manifest["state"] = "archived"
    manifest["archived_document_id"] = doc_id
    manifest["archived_at"] = report["archived_at"]
    write_yaml(stage_dir / "manifest.yaml", manifest)
    write_yaml(ROOT / "reports" / "ingestion" / f"{doc_id}-{datetime.now():%Y%m%d-%H%M%S}.yaml", report)
    print(f"Готово: {relative(destination_md)}")
    print(f"Автоматически обновлено: реестры, очередь, карта влияния и проверка целостности.")
    if warnings: print("Есть предупреждения — смотрите reports/integrity-latest.yaml")


def command_search(args: argparse.Namespace) -> None:
    print_json(search_documents(ROOT, args.query, limit=args.limit, max_chars=args.max_chars))


def command_fetch(args: argparse.Namespace) -> None:
    print_json(fetch_document(
        ROOT,
        args.document,
        clause_values=args.clauses,
        page_values=args.pages,
        query=args.query,
        context_lines=args.context_lines,
        max_chars=args.max_chars,
    ))


def command_archive_context(args: argparse.Namespace) -> None:
    print_json(archive_context(
        ROOT,
        args.stage_id,
        query=args.query,
        references=args.reference,
        max_chars=args.max_chars,
    ))


def command_route(args: argparse.Namespace) -> None:
    print_json(start_route(
        ROOT,
        task_fingerprint=args.task_id,
        complexity=args.complexity,
        ambiguity=args.ambiguity,
        criticality=args.criticality,
        context_chars=args.context_chars,
        tool_count=args.tool_count,
        side_effects=args.side_effects,
        confidence=args.confidence,
    ))


def command_route_check(args: argparse.Namespace) -> None:
    print_json(check_route(
        ROOT,
        args.run_id,
        answer_path=Path(args.answer).resolve(),
        contract_path=Path(args.contract).resolve(),
        input_path=Path(args.input).resolve() if args.input else None,
        input_tokens=args.input_tokens,
        output_tokens=args.output_tokens,
        latency_ms=args.latency_ms,
    ))


def command_route_escalate(args: argparse.Namespace) -> None:
    print_json(escalate_route(ROOT, args.run_id))


def command_route_claim(args: argparse.Namespace) -> None:
    print_json(claim_side_effects(ROOT, args.run_id, args.operation_id))


def command_routing_status(_: argparse.Namespace) -> None:
    print_json(routing_status(ROOT))


def command_routing_config(args: argparse.Namespace) -> None:
    print_json(set_routing_enabled(ROOT, args.enabled == "enable"))


def command_routing_gate(args: argparse.Namespace) -> None:
    print_json(apply_quality_gate(ROOT, Path(args.comparisons).resolve()))


def command_status(_: argparse.Namespace) -> None:
    manifest = read_yaml(META / "corpus-manifest.yaml")
    pending = [
        path.parent.name for path in (ROOT / "staging").glob("*/manifest.yaml")
        if read_yaml(path).get("state", "waiting_for_ai_analysis") == "waiting_for_ai_analysis"
    ]
    report = read_yaml(ROOT / "reports" / "integrity-latest.yaml")
    queue = read_yaml(META / "addition-queue.yaml").get("queue", [])
    priorities = {name: sum(1 for item in queue if item.get("priority") == name) for name in ("high", "medium", "low")}
    print("БАЗА ЗНАНИЙ")
    print(f"Документов: {manifest.get('document_count', 0)} | ссылок: {manifest.get('reference_count', 0)}")
    print(f"Ожидают AI-анализа: {len(pending)}" + (f" ({', '.join(pending)})" if pending else ""))
    print(f"Очередь: {len(queue)} (высокий {priorities['high']}, средний {priorities['medium']}, низкий {priorities['low']})")
    print("Последняя проверка: " + ("пройдена" if report.get("valid") else "есть ошибки/ещё не запускалась"))


def command_root(_: argparse.Namespace) -> None:
    print(ROOT)


def command_doctor(_: argparse.Namespace) -> None:
    manifest = read_yaml(META / "corpus-manifest.yaml")
    print("ПЛАГИН НВ2 — НОРМАТИВНАЯ БАЗА")
    print(f"Конфигурация: {CONFIG_PATH}")
    print(f"Корень базы: {ROOT}")
    print(f"Версия схемы: {manifest.get('schema_version')} (поддерживается {SUPPORTED_SCHEMA_VERSION})")
    print(f"Документов в индексе: {manifest.get('document_count', 0)}")
    print("Состояние: готово")


def main() -> None:
    parser = argparse.ArgumentParser(description="Конвейер нормативной базы знаний")
    commands = parser.add_subparsers(dest="command", required=True)
    stage = commands.add_parser("stage", help="Подготовить один документ для AI-анализа")
    stage.add_argument("source")
    stage.set_defaults(func=command_stage)
    intake = commands.add_parser("intake", help="Подготовить все PDF/DOCX/TXT/MD из inbox")
    intake.set_defaults(func=command_intake)
    apply = commands.add_parser("apply", help="Атомарно архивировать решение Архивария")
    apply.add_argument("decision")
    apply.add_argument("--replace", action="store_true", help="Разрешить замену существующей карточки")
    apply.set_defaults(func=command_apply)
    rebuild = commands.add_parser("rebuild-index", help="Пересобрать производные реестры")
    rebuild.set_defaults(func=command_rebuild_index)
    migrate = commands.add_parser("migrate-references", help="Вынести подробные ссылки из карточек в отдельные файлы")
    migrate.add_argument("--dry-run", action="store_true", help="Показать объём миграции без изменений")
    migrate.set_defaults(func=command_migrate_references)
    search = commands.add_parser("search", help="Найти документы и вернуть компактный JSON")
    search.add_argument("query")
    search.add_argument("--limit", type=int, default=6, help="Число результатов (1–20)")
    search.add_argument("--max-chars", type=int, default=12000, help="Предельный размер JSON (1000–50000)")
    search.add_argument("--format", choices=("json",), default="json")
    search.set_defaults(func=command_search)
    fetch = commands.add_parser("fetch", help="Извлечь точные пункты, страницы или поисковые окна")
    fetch.add_argument("document", help="ID или обозначение документа")
    fetch.add_argument("--clauses", action="append", help="Пункты через запятую; параметр можно повторять")
    fetch.add_argument("--pages", action="append", help="Страницы и диапазоны через запятую, например 3,5-7")
    fetch.add_argument("--query", help="Текст для поиска в нормализованном документе")
    fetch.add_argument("--context-lines", type=int, default=2, help="Строки контекста вокруг совпадения")
    fetch.add_argument("--max-chars", type=int, default=12000, help="Предельный объём текста фрагментов")
    fetch.add_argument("--format", choices=("json",), default="json")
    fetch.set_defaults(func=command_fetch)
    context = commands.add_parser("archive-context", help="Собрать компактный контекст для Архивария")
    context.add_argument("stage_id")
    context.add_argument("--query", help="Запрос для поиска возможных дублей")
    context.add_argument("--reference", action="append", help="Нормативная ссылка; можно повторять или перечислять через запятую")
    context.add_argument("--max-chars", type=int, default=16000, help="Предельный размер JSON (2000–50000)")
    context.add_argument("--format", choices=("json",), default="json")
    context.set_defaults(func=command_archive_context)
    route = commands.add_parser("route", help="Детерминированно выбрать worker-модель и открыть журналируемый run")
    route.add_argument("--task-id", required=True, help="Нечувствительный идентификатор задачи; в журнал попадёт только SHA-256")
    route.add_argument("--complexity", required=True, choices=("low", "medium", "high"))
    route.add_argument("--ambiguity", required=True, choices=("low", "medium", "high"))
    route.add_argument("--criticality", required=True, choices=("low", "medium", "high"))
    route.add_argument("--context-chars", required=True, type=int)
    route.add_argument("--tool-count", required=True, type=int)
    route.add_argument("--side-effects", required=True, choices=("none", "local", "external"))
    route.add_argument("--confidence", required=True, type=float)
    route.set_defaults(func=command_route)
    route_check = commands.add_parser("route-check", help="Проверить ответ и решить вопрос об однократной эскалации")
    route_check.add_argument("run_id")
    route_check.add_argument("--answer", required=True)
    route_check.add_argument("--contract", required=True)
    route_check.add_argument("--input")
    route_check.add_argument("--input-tokens", type=int)
    route_check.add_argument("--output-tokens", type=int)
    route_check.add_argument("--latency-ms", type=int)
    route_check.set_defaults(func=command_route_check)
    route_escalate = commands.add_parser("route-escalate", help="Один раз повысить уровень после провала проверки")
    route_escalate.add_argument("run_id")
    route_escalate.set_defaults(func=command_route_escalate)
    route_claim = commands.add_parser("route-claim", help="Однократно разрешить побочное действие после успешной проверки")
    route_claim.add_argument("run_id")
    route_claim.add_argument("--operation-id", required=True, help="Нечувствительный idempotency key; журналируется только SHA-256")
    route_claim.set_defaults(func=command_route_claim)
    routing_status_parser = commands.add_parser("routing-status", help="Показать feature flag и quality gate")
    routing_status_parser.set_defaults(func=command_routing_status)
    routing_config = commands.add_parser("routing-config", help="Включить или отключить каскадную маршрутизацию")
    routing_config.add_argument("enabled", choices=("enable", "disable"))
    routing_config.set_defaults(func=command_routing_config)
    routing_gate = commands.add_parser("routing-gate", help="Применить статистическое сравнение с Sol-baseline")
    routing_gate.add_argument("comparisons")
    routing_gate.set_defaults(func=command_routing_gate)
    sync = commands.add_parser("sync", help="Синхронизировать ссылки и привести в порядок очередь")
    sync.set_defaults(func=command_sync)
    validate = commands.add_parser("validate", help="Проверить структуру и связи")
    validate.set_defaults(func=command_validate)
    status = commands.add_parser("status", help="Показать короткое состояние базы")
    status.set_defaults(func=command_status)
    root = commands.add_parser("root", help="Показать путь к подключённой нормативной базе")
    root.set_defaults(func=command_root)
    doctor = commands.add_parser("doctor", help="Проверить конфигурацию плагина и совместимость базы")
    doctor.set_defaults(func=command_doctor)
    args = parser.parse_args()
    try:
        args.func(args)
    except Exception as error:
        print(f"ОШИБКА: {error}", file=sys.stderr)
        raise SystemExit(2) from error


if __name__ == "__main__":
    main()
